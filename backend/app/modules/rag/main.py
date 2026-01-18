"""
RAG module implementation with vector database and document processing
Includes comprehensive document processing, content extraction, and NLP analysis
"""
import asyncio
import io
import json
import logging
import mimetypes
import re
import time
from typing import Any, Dict, List, Optional, Tuple, Union
from datetime import datetime
from dataclasses import dataclass, asdict
from pathlib import Path
import hashlib
import base64
import numpy as np
import uuid

# Initialize logger early
logger = logging.getLogger(__name__)

# Document processing libraries (with graceful fallbacks)
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk.stem import WordNetLemmatizer

    NLTK_AVAILABLE = True
except ImportError:
    logger.warning("NLTK not available - NLP features will be limited")
    NLTK_AVAILABLE = False

try:
    import spacy

    SPACY_AVAILABLE = True
except ImportError:
    logger.warning("spaCy not available - entity extraction will be disabled")
    SPACY_AVAILABLE = False

try:
    from markitdown import MarkItDown

    MARKITDOWN_AVAILABLE = True
except ImportError:
    logger.warning("MarkItDown not available - document conversion will be limited")
    MARKITDOWN_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available - DOCX processing will be limited")
    PYTHON_DOCX_AVAILABLE = False

from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    PointStruct,
    ScoredPoint,
    Filter,
    FieldCondition,
    MatchValue,
)
from qdrant_client.http import models
import tiktoken

from app.core.config import settings
from app.core.logging import log_module_event
from app.services.base_module import BaseModule, Permission


@dataclass
class ProcessedDocument:
    """Processed document data structure"""

    id: str
    original_filename: str
    file_type: str
    mime_type: str
    content: str
    extracted_text: str
    metadata: Dict[str, Any]
    word_count: int
    sentence_count: int
    language: str
    entities: List[Dict[str, Any]]
    keywords: List[str]
    processing_time: float
    processed_at: datetime
    file_hash: str
    file_size: int
    embedding: Optional[List[float]] = None
    source_url: Optional[str] = None
    created_at: datetime = None

    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.utcnow()


@dataclass
class ContentValidationResult:
    """Content validation result"""

    is_valid: bool
    issues: List[str]
    security_score: float
    content_type: str
    language_confidence: float


# Keep Document class for backward compatibility
@dataclass
class Document:
    """Simple document data structure for backward compatibility"""

    id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None
    created_at: datetime = None

    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.utcnow()


@dataclass
class SearchResult:
    """Search result data structure"""

    document: Document
    score: float
    relevance_score: float


class RAGModule(BaseModule):
    """RAG module for document storage, retrieval, and augmented generation with integrated content processing"""

    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(module_id="rag", config=config)
        self.enabled = False
        self.qdrant_client: Optional[QdrantClient] = None
        self.default_collection_name = "documents"  # Keep for backward compatibility
        self.embedding_model = None
        self.embedding_service = None
        self.tokenizer = None

        # Set improved default configuration
        self.config = {
            "chunk_size": 300,  # Reduced from 400 for better precision
            "chunk_overlap": 50,  # Added overlap for context preservation
            "max_results": 10,
            "score_threshold": 0.3,  # Increased from 0.0 to filter low-quality results
            "enable_hybrid": True,  # Enable hybrid search (vector + BM25)
            "hybrid_weights": {"vector": 0.7, "bm25": 0.3},  # Weight for hybrid scoring
        }
        # Update with any provided config
        if config:
            self.config.update(config)

        # Ensure embedding model configured (defaults to local BGE-small-en)
        default_embedding_model = getattr(
            settings, "RAG_EMBEDDING_MODEL", "BAAI/bge-small-en-v1.5"
        )
        self.config.setdefault("embedding_model", default_embedding_model)
        self.default_embedding_model = default_embedding_model

        # Content processing components
        self.nlp_model = None
        self.lemmatizer = None
        self.stop_words = set()
        self.markitdown = None
        self.supported_types = {
            "text/plain": self._process_text,
            "application/pdf": self._process_with_markitdown,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._process_docx,
            "application/msword": self._process_docx,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": self._process_with_markitdown,
            "application/vnd.ms-excel": self._process_with_markitdown,
            "text/html": self._process_html,
            "application/json": self._process_json,
            "application/x-ndjson": self._process_jsonl,  # JSONL support
            "text/markdown": self._process_markdown,
            "text/csv": self._process_csv,
        }

        self.stats = {
            "documents_indexed": 0,
            "documents_processed": 0,
            "total_processing_time": 0,
            "average_processing_time": 0,
            "searches_performed": 0,
            "average_search_time": 0.0,
            "cache_hits": 0,
            "errors": 0,
            "supported_types": len(self.supported_types),
        }
        self.search_cache = {}
        self.collection_vector_sizes: Dict[str, int] = {}

    def get_required_permissions(self) -> List[Permission]:
        """Return list of permissions this module requires"""
        return [
            Permission("documents", "index", "Index new documents"),
            Permission("documents", "search", "Search documents"),
            Permission("documents", "delete", "Delete documents"),
            Permission("collections", "manage", "Manage collections"),
            Permission("settings", "configure", "Configure RAG settings"),
        ]

    async def initialize(self):
        """Initialize the RAG module with content processing capabilities"""

        try:
            # Initialize Qdrant client (settings from core.config)
            qdrant_url = f"http://{settings.QDRANT_HOST}:{settings.QDRANT_PORT}"
            self.qdrant_client = QdrantClient(url=qdrant_url)

            # Initialize tokenizer
            self.tokenizer = tiktoken.get_encoding("cl100k_base")

            # Initialize embedding model
            self.embedding_model = await self._initialize_embedding_model()

            # Initialize content processing components
            await self._initialize_content_processing()

            # Create default collection if it doesn't exist
            await self._ensure_collection_exists(self.default_collection_name)

            self.enabled = True
            self.initialized = True
            log_module_event(
                "rag",
                "initialized",
                {
                    "vector_db": self.config.get("vector_db", "qdrant"),
                    "embedding_model": self.embedding_model.get(
                        "model_name", self.default_embedding_model
                    ),
                    "chunk_size": self.config.get("chunk_size", 400),
                    "max_results": self.config.get("max_results", 10),
                    "supported_file_types": list(self.supported_types.keys()),
                    "nltk_ready": True,
                    "spacy_ready": self.nlp_model is not None,
                    "markitdown_ready": self.markitdown is not None,
                },
            )

        except Exception as e:
            logger.error(f"Failed to initialize RAG module: {e}")
            log_module_event("rag", "initialization_failed", {"error": str(e)})
            self.enabled = False
            raise

    def _generate_file_hash(self, content: bytes) -> str:
        """Generate SHA-256 hash of file content"""
        return hashlib.sha256(content).hexdigest()

    def _detect_mime_type(self, filename: str, content: bytes) -> str:
        """Detect MIME type of file"""
        # Try to detect from filename
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            return mime_type

        # Check for JSONL file extension
        if filename.lower().endswith(".jsonl"):
            return "application/x-ndjson"

        # Try to detect from content
        if content.startswith(b"%PDF"):
            return "application/pdf"
        elif content.startswith(b"PK"):
            # This could be DOCX, XLSX, or other Office formats
            if filename.lower().endswith((".docx", ".docm")):
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif filename.lower().endswith((".xlsx", ".xlsm")):
                return (
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            else:
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif content.startswith(b"\xd0\xcf\x11\xe0"):
            # Old Office format (DOC, XLS)
            if filename.lower().endswith(".xls"):
                return "application/vnd.ms-excel"
            else:
                return "application/msword"
        elif content.startswith(b"<html") or content.startswith(b"<!DOCTYPE"):
            return "text/html"
        elif content.startswith(b"{") or content.startswith(b"["):
            # Check if it's JSONL by looking for newline-delimited JSON
            try:
                content_str = content.decode("utf-8", errors="ignore")
                lines = content_str.split("\n")
                # Filter out empty lines
                non_empty_lines = [line.strip() for line in lines[:10] if line.strip()]

                # If we have multiple non-empty lines that all start with {, it's likely JSONL
                if len(non_empty_lines) > 1 and all(
                    line.startswith("{") and line.endswith("}") for line in non_empty_lines[:5]
                ):
                    # Additional validation: try parsing a few lines as JSON
                    import json
                    valid_json_lines = 0
                    for line in non_empty_lines[:3]:
                        try:
                            json.loads(line)
                            valid_json_lines += 1
                        except:
                            break

                    if valid_json_lines > 1:
                        return "application/x-ndjson"
            except:
                pass
            return "application/json"
        else:
            return "text/plain"

    def _detect_language(self, text: str) -> Tuple[str, float]:
        """Detect language of text (simplified implementation)"""
        if len(text) < 50:
            return "unknown", 0.0

        # Simple heuristic based on common English words
        english_words = {
            "the",
            "and",
            "or",
            "but",
            "in",
            "on",
            "at",
            "to",
            "for",
            "of",
            "with",
            "by",
            "is",
            "are",
            "was",
            "were",
            "be",
            "been",
            "being",
            "have",
            "has",
            "had",
            "do",
            "does",
            "did",
            "will",
            "would",
            "could",
            "should",
            "may",
            "might",
            "must",
            "can",
            "shall",
        }

        if NLTK_AVAILABLE:
            words = word_tokenize(text.lower())
        else:
            # Fallback to simple whitespace tokenization
            words = text.lower().split()

        english_count = sum(1 for word in words if word in english_words)
        confidence = min(english_count / len(words), 1.0) if words else 0.0

        return "en" if confidence > 0.1 else "unknown", confidence

    def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities from text"""
        if not self.nlp_model:
            return []

        try:
            doc = self.nlp_model(text[:10000])  # Limit text length for performance
            entities = []

            for ent in doc.ents:
                entities.append(
                    {
                        "text": ent.text,
                        "label": ent.label_,
                        "start": ent.start_char,
                        "end": ent.end_char,
                        "confidence": float(ent._.get("score", 0.0))
                        if hasattr(ent._, "score")
                        else 0.0,
                    }
                )

            return entities

        except Exception as e:
            logger.error(f"Error extracting entities: {e}")
            return []

    def _extract_keywords(self, text: str, max_keywords: int = 20) -> List[str]:
        """Extract keywords from text"""
        try:
            if NLTK_AVAILABLE:
                words = word_tokenize(text.lower())
            else:
                # Fallback to simple whitespace tokenization
                words = text.lower().split()

            words = [
                word for word in words if word.isalpha() and word not in self.stop_words
            ]

            if self.lemmatizer and NLTK_AVAILABLE:
                words = [self.lemmatizer.lemmatize(word) for word in words]

            # Simple frequency-based keyword extraction
            word_freq = {}
            for word in words:
                word_freq[word] = word_freq.get(word, 0) + 1

            # Sort by frequency and return top keywords
            keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            return [word for word, freq in keywords[:max_keywords] if freq > 1]

        except Exception as e:
            logger.error(f"Error extracting keywords: {e}")
            return []

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove control characters except newlines and tabs
        text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

        # Normalize quotes
        text = re.sub(r'[""' "`]", '"', text)

        # Remove excessive punctuation
        text = re.sub(r"[.]{3,}", "...", text)
        text = re.sub(r"[!]{2,}", "!", text)
        text = re.sub(r"[?]{2,}", "?", text)

        return text.strip()

    def _validate_content(
        self, content: str, file_type: str
    ) -> ContentValidationResult:
        """Validate and score content for security and quality"""
        issues = []
        security_score = 100.0

        # Check for potentially malicious content
        if "<script" in content.lower() or "javascript:" in content.lower():
            issues.append("Potentially malicious JavaScript content detected")
            security_score -= 30

        if re.search(r"<iframe|<object|<embed", content, re.IGNORECASE):
            issues.append("Embedded content detected")
            security_score -= 20

        # Check for suspicious URLs
        if re.search(
            r"https?://[^\s]+\.(exe|bat|cmd|scr|vbs|js)", content, re.IGNORECASE
        ):
            issues.append("Suspicious executable URLs detected")
            security_score -= 40

        # Check content length
        if len(content) > 1000000:  # 1MB limit
            issues.append("Content exceeds maximum size limit")
            security_score -= 10

        # Detect language
        language, lang_confidence = self._detect_language(content)

        return ContentValidationResult(
            is_valid=len(issues) == 0,
            issues=issues,
            security_score=max(0, security_score),
            content_type=file_type,
            language_confidence=lang_confidence,
        )

    async def cleanup(self):
        """Cleanup RAG resources"""
        if self.qdrant_client:
            self.qdrant_client.close()
            self.qdrant_client = None

        if self.embedding_service:
            await self.embedding_service.cleanup()
            self.embedding_service = None

        # Cleanup content processing resources
        self.nlp_model = None
        self.lemmatizer = None
        self.markitdown = None
        self.stop_words.clear()

        self.enabled = False
        self.search_cache.clear()
        log_module_event("rag", "cleanup", {"success": True})

    async def _initialize_embedding_model(self):
        """Initialize embedding model"""
        # Prefer enhanced embedding service (rate limiting + retry)
        from app.services.enhanced_embedding_service import (
            enhanced_embedding_service as embedding_service,
        )

        model_name = self.config.get("embedding_model", self.default_embedding_model)
        embedding_service.model_name = model_name

        # Initialize the embedding service
        success = await embedding_service.initialize()

        if success:
            self.embedding_service = embedding_service
            logger.info(f"Successfully initialized embedding service with {model_name}")
            return {
                "model_name": model_name,
                "dimension": embedding_service.dimension or 384,
            }
        else:
            # Fallback to mock implementation
            logger.warning("Failed to initialize embedding model, using fallback")
            self.embedding_service = None
            return {
                "model_name": model_name,
                "dimension": 384,  # Default dimension matching local bge-small embeddings
            }

    async def _initialize_content_processing(self):
        """Initialize content processing components"""
        try:
            # Download required NLTK data
            await self._download_nltk_data()

            # Initialize NLP components
            if NLTK_AVAILABLE:
                self.lemmatizer = WordNetLemmatizer()
                self.stop_words = set(stopwords.words("english"))
            else:
                self.lemmatizer = None
                self.stop_words = set()

            # Initialize spaCy model
            await self._initialize_spacy_model()

            # Initialize MarkItDown
            if MARKITDOWN_AVAILABLE:
                self.markitdown = MarkItDown()
            else:
                self.markitdown = None

        except Exception as e:
            logger.warning(
                f"Failed to initialize some content processing components: {e}"
            )

    async def _download_nltk_data(self):
        """Download required NLTK data"""
        if not NLTK_AVAILABLE:
            return

        try:
            nltk.download("punkt", quiet=True)
            nltk.download("stopwords", quiet=True)
            nltk.download("wordnet", quiet=True)
            nltk.download("averaged_perceptron_tagger", quiet=True)
            nltk.download("omw-1.4", quiet=True)
        except Exception as e:
            logger.warning(f"Failed to download NLTK data: {e}")

    async def _initialize_spacy_model(self):
        """Initialize spaCy model for NLP tasks"""
        if not SPACY_AVAILABLE:
            self.nlp_model = None
            return

        try:
            self.nlp_model = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning(
                "spaCy model 'en_core_web_sm' not found. NLP features will be limited."
            )
            self.nlp_model = None

    async def _get_collections_safely(self) -> List[str]:
        """Get list of collections using raw HTTP to avoid Pydantic validation issues"""
        try:
            import httpx

            qdrant_host = getattr(settings, "QDRANT_HOST", "localhost")
            qdrant_port = getattr(settings, "QDRANT_PORT", 6333)
            qdrant_url = f"http://{qdrant_host}:{qdrant_port}"

            async with httpx.AsyncClient() as client:
                response = await client.get(f"{qdrant_url}/collections")
                if response.status_code == 200:
                    data = response.json()
                    result = data.get("result", {})
                    collections = result.get("collections", [])
                    return [
                        col.get("name", "") for col in collections if col.get("name")
                    ]
                else:
                    logger.warning(
                        f"Failed to get collections via HTTP: {response.status_code}"
                    )
                    return []
        except Exception as e:
            logger.error(f"Error getting collections safely: {e}")
            # Fallback to direct client call with error handling
            try:
                collections = self.qdrant_client.get_collections()
                return [col.name for col in collections.collections]
            except Exception as fallback_error:
                logger.error(f"Fallback collection fetch also failed: {fallback_error}")
                return []

    async def _get_collection_info_safely(self, collection_name: str) -> Dict[str, Any]:
        """Get collection information using raw HTTP to avoid Pydantic validation issues"""
        try:
            import httpx

            qdrant_host = getattr(settings, "QDRANT_HOST", "localhost")
            qdrant_port = getattr(settings, "QDRANT_PORT", 6333)
            qdrant_url = f"http://{qdrant_host}:{qdrant_port}"

            async with httpx.AsyncClient() as client:
                response = await client.get(
                    f"{qdrant_url}/collections/{collection_name}"
                )
                if response.status_code == 200:
                    data = response.json()
                    result = data.get("result", {})

                    # Extract relevant information safely
                    collection_info = {
                        "points_count": result.get("points_count", 0),
                        "status": result.get("status", "unknown"),
                        "vector_size": 384,  # Default fallback
                    }

                    # Try to get vector dimension from config
                    try:
                        config = result.get("config", {})
                        params = config.get("params", {})
                        vectors = params.get("vectors", {})

                        if isinstance(vectors, dict) and "size" in vectors:
                            collection_info["vector_size"] = vectors["size"]
                        elif isinstance(vectors, dict):
                            # Handle named vectors or default vector
                            if "default" in vectors:
                                collection_info["vector_size"] = vectors["default"].get(
                                    "size", 384
                                )
                            else:
                                # Take first vector config if no default
                                first_vector = next(iter(vectors.values()), {})
                                collection_info["vector_size"] = first_vector.get(
                                    "size", 384
                                )
                    except Exception:
                        # Keep default fallback
                        pass

                    return collection_info
                else:
                    logger.warning(
                        f"Failed to get collection info via HTTP: {response.status_code}"
                    )
                    return {"points_count": 0, "status": "error", "vector_size": 384}
        except Exception as e:
            logger.error(f"Error getting collection info safely: {e}")
            return {"points_count": 0, "status": "error", "vector_size": 384}

    async def _ensure_collection_exists(self, collection_name: str = None):
        """Ensure the specified collection exists"""
        collection_name = collection_name or self.default_collection_name

        try:
            # Use safe collection fetching to avoid Pydantic validation errors
            collection_names = await self._get_collections_safely()

            if collection_name not in collection_names:
                # Create collection with the current embedding dimension
                vector_dimension = self.embedding_model.get(
                    "dimension",
                    getattr(self.embedding_service, "dimension", 384) or 384,
                )

                self.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(
                        size=vector_dimension, distance=Distance.COSINE
                    ),
                )
                self.collection_vector_sizes[collection_name] = vector_dimension
                log_module_event(
                    "rag", "collection_created", {"collection": collection_name}
                )
            else:
                # Cache existing collection vector size for later alignment
                try:
                    info = self.qdrant_client.get_collection(collection_name)
                    vectors_param = (
                        getattr(info.config.params, "vectors", None)
                        if hasattr(info, "config")
                        else None
                    )
                    existing_size = None
                    if vectors_param is not None and hasattr(vectors_param, "size"):
                        existing_size = vectors_param.size
                    elif isinstance(vectors_param, dict):
                        existing_size = vectors_param.get("size")

                    if existing_size:
                        self.collection_vector_sizes[collection_name] = existing_size
                except Exception as inner_error:
                    logger.debug(
                        f"Unable to cache vector size for collection {collection_name}: {inner_error}"
                    )

        except Exception as e:
            logger.error(f"Error ensuring collection exists: {e}")
            raise

    async def create_collection(self, collection_name: str) -> bool:
        """Create a new Qdrant collection"""
        try:
            await self._ensure_collection_exists(collection_name)
            return True
        except Exception as e:
            logger.error(f"Error creating collection {collection_name}: {e}")
            return False

    async def delete_collection(self, collection_name: str) -> bool:
        """Delete a Qdrant collection"""
        try:
            # Use safe collection fetching to avoid Pydantic validation errors
            collection_names = await self._get_collections_safely()

            if collection_name in collection_names:
                self.qdrant_client.delete_collection(collection_name)
                log_module_event(
                    "rag", "collection_deleted", {"collection": collection_name}
                )
                return True
            else:
                logger.warning(f"Collection {collection_name} does not exist")
                return False

        except Exception as e:
            logger.error(f"Error deleting collection {collection_name}: {e}")
            return False

    async def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text"""
        if self.embedding_service:
            vector = await self.embedding_service.get_embedding(text)
            return vector
        else:
            # Fallback to deterministic random embedding for consistency
            np.random.seed(hash(text) % 2**32)
            fallback_dim = self.embedding_model.get(
                "dimension", getattr(self.embedding_service, "dimension", 384) or 384
            )
            return np.random.random(fallback_dim).tolist()

    async def _generate_embeddings(
        self, texts: List[str], is_document: bool = True
    ) -> List[List[float]]:
        """Generate embeddings for multiple texts (batch processing)"""
        if self.embedding_service:
            # Add task-specific prefixes for better E5 model performance
            if is_document:
                # For document passages, use "passage:" prefix
                prefixed_texts = [f"passage: {text}" for text in texts]
            else:
                # For queries, use "query:" prefix (handled in search method)
                prefixed_texts = texts

            # Use real embedding service for batch processing
            backend = getattr(self.embedding_service, "backend", "unknown")
            start_time = time.time()
            logger.info(
                "Embedding batch requested",
                extra={
                    "backend": backend,
                    "model": getattr(self.embedding_service, "model_name", "unknown"),
                    "count": len(prefixed_texts),
                    "scope": "documents" if is_document else "queries",
                },
            )
            embeddings = await self.embedding_service.get_embeddings(prefixed_texts)
            duration = time.time() - start_time
            logger.info(
                "Embedding batch finished",
                extra={
                    "backend": backend,
                    "model": getattr(self.embedding_service, "model_name", "unknown"),
                    "count": len(embeddings),
                    "scope": "documents" if is_document else "queries",
                    "duration_sec": round(duration, 4),
                },
            )
            return embeddings
        else:
            # Fallback to individual processing
            logger.warning(
                "Embedding service unavailable, falling back to per-item generation",
                extra={
                    "count": len(texts),
                    "scope": "documents" if is_document else "queries",
                },
            )
            embeddings = []
            for text in texts:
                embedding = await self._generate_embedding(text)
                embeddings.append(embedding)
            return embeddings

    def _get_collection_vector_size(self, collection_name: Optional[str]) -> int:
        """Return the expected vector size for a collection, caching results."""
        default_dim = self.embedding_model.get(
            "dimension", getattr(self.embedding_service, "dimension", 384) or 384
        )

        if not collection_name:
            return default_dim

        if collection_name in self.collection_vector_sizes:
            return self.collection_vector_sizes[collection_name]

        try:
            info = self.qdrant_client.get_collection(collection_name)
            vectors_param = (
                getattr(info.config.params, "vectors", None)
                if hasattr(info, "config")
                else None
            )
            existing_size = None
            if vectors_param is not None and hasattr(vectors_param, "size"):
                existing_size = vectors_param.size
            elif isinstance(vectors_param, dict):
                existing_size = vectors_param.get("size")

            if existing_size:
                self.collection_vector_sizes[collection_name] = existing_size
                return existing_size
        except Exception as e:
            logger.debug(f"Unable to determine vector size for {collection_name}: {e}")

        self.collection_vector_sizes[collection_name] = default_dim
        return default_dim

    def _align_embedding_dimension(
        self, vector: List[float], collection_name: Optional[str]
    ) -> List[float]:
        """Pad or truncate embeddings to match the target collection dimension."""
        if vector is None:
            return vector

        target_dim = self._get_collection_vector_size(collection_name)
        current_dim = len(vector)

        if current_dim == target_dim:
            return vector
        if current_dim > target_dim:
            return vector[:target_dim]
        padding = [0.0] * (target_dim - current_dim)
        return vector + padding

    def _chunk_text(self, text: str, chunk_size: int = None) -> List[str]:
        """Split text into overlapping chunks for better context preservation"""
        chunk_size = chunk_size or self.config.get("chunk_size", 300)
        chunk_overlap = self.config.get("chunk_overlap", 50)

        # Tokenize text
        tokens = self.tokenizer.encode(text)

        # Split into chunks with overlap
        chunks = []
        start_idx = 0

        while start_idx < len(tokens):
            end_idx = min(start_idx + chunk_size, len(tokens))
            chunk_tokens = tokens[start_idx:end_idx]
            chunk_text = self.tokenizer.decode(chunk_tokens)

            # Only add non-empty chunks
            if chunk_text.strip():
                chunks.append(chunk_text)

            # Move to next chunk with overlap
            # Ensure we make progress and don't loop infinitely
            start_idx += chunk_size - chunk_overlap
            if start_idx >= len(tokens):
                break

            # Safety check to prevent infinite loop
            if start_idx <= end_idx - chunk_size:
                start_idx = end_idx

        return chunks

    async def _process_text(self, content: bytes, filename: str) -> str:
        """Process plain text files"""
        try:
            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            # Fallback to utf-8 with error handling
            return content.decode("utf-8", errors="replace")

        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            return ""

    async def _process_with_markitdown(self, content: bytes, filename: str) -> str:
        """Process documents using MarkItDown (PDF, DOCX, DOC, XLSX, XLS)"""
        try:
            if not self.markitdown:
                raise RuntimeError("MarkItDown not initialized")

            # Create a temporary file path for the content
            import tempfile
            import os

            # Get file extension from filename
            file_ext = Path(filename).suffix.lower()
            if not file_ext:
                # Try to determine extension from mime type
                mime_type = self._detect_mime_type(filename, content)
                if mime_type == "application/pdf":
                    file_ext = ".pdf"
                elif mime_type in [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ]:
                    file_ext = ".docx"
                elif mime_type == "application/msword":
                    file_ext = ".doc"
                elif (
                    mime_type
                    == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                ):
                    file_ext = ".xlsx"
                elif mime_type == "application/vnd.ms-excel":
                    file_ext = ".xls"
                else:
                    file_ext = ".bin"

            # Write content to temporary file
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=file_ext
            ) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name

            try:
                # Convert document to markdown using MarkItDown in a thread pool to avoid blocking
                import concurrent.futures
                import asyncio

                logger.info(f"Starting MarkItDown conversion for {filename}")

                def convert_sync():
                    """Synchronous conversion function to run in thread pool"""
                    return self.markitdown.convert(temp_path)

                # Run the synchronous conversion in a thread pool with timeout
                loop = asyncio.get_event_loop()
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    try:
                        result = await asyncio.wait_for(
                            loop.run_in_executor(executor, convert_sync),
                            timeout=120.0,  # 2 minute timeout for MarkItDown conversion
                        )
                    except asyncio.TimeoutError:
                        logger.error(f"MarkItDown conversion timed out for {filename}")
                        raise RuntimeError(
                            f"Document conversion timed out after 2 minutes for {filename}"
                        )

                if result and hasattr(result, "text_content"):
                    converted_text = result.text_content
                elif result and isinstance(result, str):
                    converted_text = result
                else:
                    # Fallback if result format is unexpected
                    converted_text = str(result) if result else ""

                logger.info(
                    f"Successfully converted {filename} using MarkItDown ({len(converted_text)} characters)"
                )
                return converted_text

            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass

        except Exception as e:
            logger.error(f"Error processing {filename} with MarkItDown: {e}")
            # Fallback to basic text extraction attempt
            try:
                return content.decode("utf-8", errors="replace")
            except:
                return f"Error processing {filename}: {str(e)}"

    async def _process_docx(self, content: bytes, filename: str) -> str:
        """Process DOCX files using python-docx (more reliable than MarkItDown)"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                logger.warning(
                    f"python-docx not available, falling back to MarkItDown for {filename}"
                )
                return await self._process_with_markitdown(content, filename)

            # Create a temporary file for python-docx processing
            import tempfile
            import os

            logger.info(f"Starting DOCX processing for {filename} using python-docx")

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name

            try:
                # Process in a thread pool to avoid blocking
                import concurrent.futures
                import asyncio

                def extract_docx_text():
                    """Extract text from DOCX file synchronously"""
                    doc = DocxDocument(temp_path)
                    text_parts = []

                    # Extract paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())

                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_parts.append(" | ".join(row_text))

                    return "\n\n".join(text_parts)

                # Run extraction in thread pool with timeout
                loop = asyncio.get_event_loop()
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    try:
                        extracted_text = await asyncio.wait_for(
                            loop.run_in_executor(executor, extract_docx_text),
                            timeout=30.0,  # 30 second timeout for DOCX processing
                        )
                    except asyncio.TimeoutError:
                        logger.error(f"DOCX processing timed out for {filename}")
                        raise RuntimeError(
                            f"DOCX processing timed out after 30 seconds for {filename}"
                        )

                logger.info(
                    f"Successfully processed {filename} using python-docx ({len(extracted_text)} characters)"
                )
                return extracted_text

            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass

        except Exception as e:
            logger.error(f"Error processing DOCX file {filename}: {e}")
            # Fallback to MarkItDown if python-docx fails
            try:
                logger.info(f"Falling back to MarkItDown for {filename}")
                return await self._process_with_markitdown(content, filename)
            except Exception as fallback_error:
                logger.error(
                    f"Both python-docx and MarkItDown failed for {filename}: {fallback_error}"
                )
                return f"Error processing DOCX {filename}: {str(e)}"

    async def _process_html(self, content: bytes, filename: str) -> str:
        """Process HTML files"""
        try:
            html_content = content.decode("utf-8", errors="replace")
            # Simple HTML tag removal
            text = re.sub(r"<[^>]+>", "", html_content)
            # Decode HTML entities
            text = (
                text.replace("&amp;", "&")
                .replace("&lt;", "<")
                .replace("&gt;", ">")
                .replace("&quot;", '"')
                .replace("&#39;", "'")
            )
            return text

        except Exception as e:
            logger.error(f"Error processing HTML file: {e}")
            return ""

    async def _process_json(self, content: bytes, filename: str) -> str:
        """Process JSON files"""
        try:
            json_str = content.decode("utf-8", errors="ignore")
            json_data = json.loads(json_str)
            # Convert JSON to readable text
            return json.dumps(json_data, indent=2)

        except json.JSONDecodeError as e:
            # Check if this might be JSONL content that was misdetected
            try:
                lines = json_str.split("\n")
                # Filter out empty lines
                non_empty_lines = [line.strip() for line in lines if line.strip()]

                # If multiple valid JSON lines, treat as JSONL
                if len(non_empty_lines) > 1:
                    logger.warning(f"File '{filename}' appears to be JSONL format, processing as JSONL")
                    # Call JSONL processor directly
                    return await self._process_jsonl(content, filename)

                logger.error(f"Error processing JSON file '{filename}': {e}")
                return ""
            except Exception as fallback_e:
                logger.error(f"Error processing JSON file '{filename}': {e}, fallback also failed: {fallback_e}")
                return ""
        except Exception as e:
            logger.error(f"Error processing JSON file '{filename}': {e}")
            return ""

    async def _process_markdown(self, content: bytes, filename: str) -> str:
        """Process Markdown files"""
        try:
            md_content = content.decode("utf-8", errors="replace")
            # Simple markdown processing - remove formatting
            text = re.sub(r"#+\s*", "", md_content)  # Remove headers
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # Bold
            text = re.sub(r"\*(.+?)\*", r"\1", text)  # Italic
            text = re.sub(r"`(.+?)`", r"\1", text)  # Code
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)  # Links
            return text

        except Exception as e:
            logger.error(f"Error processing Markdown file: {e}")
            return ""

    async def _process_csv(self, content: bytes, filename: str) -> str:
        """Process CSV files"""
        try:
            csv_content = content.decode("utf-8", errors="replace")
            # Convert CSV to readable text
            lines = csv_content.split("\n")
            processed_lines = []

            for line in lines[:100]:  # Limit to first 100 lines
                if line.strip():
                    processed_lines.append(line.replace(",", " | "))

            return "\n".join(processed_lines)

        except Exception as e:
            logger.error(f"Error processing CSV file: {e}")
            return ""

    async def _process_jsonl(self, content: bytes, filename: str) -> str:
        """Process JSONL files (newline-delimited JSON)

        Specifically optimized for helpjuice-export.jsonl format:
        - Each line contains a JSON object with 'id' and 'payload'
        - Payload contains 'question', 'language', and 'answer' fields
        - Combines question and answer into searchable content

        Performance optimizations:
        - Processes articles in smaller batches to reduce memory usage
        - Uses streaming approach for large files
        """
        try:
            # Use streaming approach for large files
            jsonl_content = content.decode("utf-8", errors="replace")
            lines = jsonl_content.strip().split("\n")

            processed_articles = []
            batch_size = 50  # Process in batches of 50 articles

            for line_num, line in enumerate(lines, 1):
                if not line.strip():
                    continue

                try:
                    # Parse each JSON line
                    data = json.loads(line)

                    # Handle helpjuice export format
                    if "payload" in data:
                        payload = data["payload"]
                        article_id = data.get("id", f"article_{line_num}")

                        # Extract fields
                        question = payload.get("question", "")
                        answer = payload.get("answer", "")
                        language = payload.get("language", "EN")

                        # Combine question and answer for better search
                        if question or answer:
                            # Format as Q&A for better context
                            article_text = f"## {question}\n\n{answer}\n\n"

                            # Add language tag if not English
                            if language != "EN":
                                article_text = f"[{language}] {article_text}"

                            # Add metadata separator
                            article_text += f"---\nArticle ID: {article_id}\nLanguage: {language}\n\n"

                            processed_articles.append(article_text)

                    # Handle generic JSONL format
                    else:
                        # Convert the entire JSON object to readable text
                        json_text = json.dumps(data, indent=2, ensure_ascii=False)
                        processed_articles.append(json_text + "\n\n")

                except json.JSONDecodeError as e:
                    logger.warning(f"Error parsing JSONL line {line_num}: {e}")
                    continue
                except Exception as e:
                    logger.warning(f"Error processing JSONL line {line_num}: {e}")
                    continue

            # Combine all articles
            combined_text = "\n".join(processed_articles)

            logger.info(
                f"Successfully processed {len(processed_articles)} articles from JSONL file {filename}"
            )
            return combined_text

        except Exception as e:
            logger.error(f"Error processing JSONL file {filename}: {e}")
            return ""

    def _generate_document_id(self, content: str, metadata: Dict[str, Any]) -> str:
        """Generate unique document ID"""
        content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
        metadata_hash = hashlib.sha256(
            json.dumps(metadata, sort_keys=True).encode()
        ).hexdigest()[:8]
        return f"{content_hash}_{metadata_hash}"

    async def process_document(
        self, file_data: bytes, filename: str, metadata: Dict[str, Any] = None
    ) -> ProcessedDocument:
        """Process a document and extract content"""
        if not self.enabled:
            raise RuntimeError("RAG module not initialized")

        import time

        start_time = time.time()

        try:
            logger.info(f"Starting document processing pipeline for {filename}")

            # Generate file hash and ID
            file_hash = self._generate_file_hash(file_data)
            doc_id = f"{file_hash}_{int(time.time())}"
            logger.info(f"Generated document ID: {doc_id}")

            # Detect MIME type
            mime_type = self._detect_mime_type(filename, file_data)
            # Special handling for JSONL files - use extension instead of MIME family
            if mime_type == "application/x-ndjson" or filename.lower().endswith('.jsonl'):
                file_type = "jsonl"
            else:
                file_type = mime_type.split("/")[0]
            logger.info(f"Detected MIME type: {mime_type}, file type: {file_type}")

            # Check if file type is supported
            if mime_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {mime_type}")

            # Extract content using appropriate processor
            processor = self.supported_types[mime_type]
            logger.info(f"Using processor: {processor.__name__} for {filename}")
            extracted_text = await processor(file_data, filename)
            logger.info(
                f"Content extraction completed for {filename}, extracted {len(extracted_text)} characters"
            )

            # Clean the extracted text
            logger.info(f"Starting text cleaning for {filename}")
            cleaned_text = self._clean_text(extracted_text)
            logger.info(
                f"Text cleaning completed for {filename}, final text length: {len(cleaned_text)}"
            )

            # Validate content
            logger.info(f"Starting content validation for {filename}")
            validation_result = self._validate_content(cleaned_text, file_type)
            logger.info(f"Content validation completed for {filename}")

            if not validation_result.is_valid:
                logger.warning(f"Content validation issues: {validation_result.issues}")

            # Extract linguistic features
            logger.info(f"Starting linguistic analysis for {filename}")
            if NLTK_AVAILABLE and cleaned_text:
                logger.info(f"Using NLTK for tokenization of {filename}")
                sentences = sent_tokenize(cleaned_text)
                words = word_tokenize(cleaned_text)
            elif cleaned_text:
                logger.info(f"Using fallback tokenization for {filename}")
                # Fallback to simple tokenization
                sentences = cleaned_text.split(".")
                words = cleaned_text.split()
            else:
                logger.warning(f"No text content for linguistic analysis in {filename}")
                sentences = []
                words = []

            logger.info(
                f"Tokenization completed for {filename}: {len(sentences)} sentences, {len(words)} words"
            )

            # Detect language
            logger.info(f"Starting language detection for {filename}")
            language, lang_confidence = self._detect_language(cleaned_text)
            logger.info(
                f"Language detection completed for {filename}: {language} (confidence: {lang_confidence:.2f})"
            )

            # Extract entities and keywords
            logger.info(f"Starting entity extraction for {filename}")
            entities = self._extract_entities(cleaned_text)
            logger.info(
                f"Entity extraction completed for {filename}: found {len(entities)} entities"
            )

            logger.info(f"Starting keyword extraction for {filename}")
            keywords = self._extract_keywords(cleaned_text)
            logger.info(
                f"Keyword extraction completed for {filename}: found {len(keywords)} keywords"
            )

            # Calculate processing time
            processing_time = time.time() - start_time

            # Create processed document
            logger.info(f"Creating ProcessedDocument object for {filename}")
            processed_doc = ProcessedDocument(
                id=doc_id,
                original_filename=filename,
                file_type=file_type,
                mime_type=mime_type,
                content=cleaned_text,
                extracted_text=extracted_text,
                metadata={
                    **(metadata or {}),
                    "validation": asdict(validation_result),
                    "file_size": len(file_data),
                    "processing_stats": {
                        "processing_time": processing_time,
                        "processor_used": processor.__name__,
                    },
                },
                word_count=len(words),
                sentence_count=len(sentences),
                language=language,
                entities=entities,
                keywords=keywords,
                processing_time=processing_time,
                processed_at=datetime.utcnow(),
                file_hash=file_hash,
                file_size=len(file_data),
            )
            logger.info(f"ProcessedDocument created for {filename}")

            # Update stats
            self.stats["documents_processed"] += 1
            self.stats["total_processing_time"] += processing_time
            self.stats["average_processing_time"] = (
                self.stats["total_processing_time"] / self.stats["documents_processed"]
            )

            log_module_event(
                "rag",
                "document_processed",
                {
                    "document_id": doc_id,
                    "filename": filename,
                    "file_type": file_type,
                    "word_count": len(words),
                    "processing_time": processing_time,
                    "language": language,
                    "entities_count": len(entities),
                    "keywords_count": len(keywords),
                },
            )

            logger.info(
                f"Document processing completed successfully for {filename} in {processing_time:.2f} seconds"
            )
            return processed_doc

        except Exception as e:
            self.stats["errors"] += 1
            logger.error(f"Error processing document {filename}: {e}")
            log_module_event(
                "rag", "processing_failed", {"filename": filename, "error": str(e)}
            )
            raise

    async def index_document(
        self, content: str, metadata: Dict[str, Any] = None, collection_name: str = None
    ) -> str:
        """Index a document in the vector database (backward compatibility method)"""
        if not self.enabled:
            raise RuntimeError("RAG module not initialized")

        collection_name = collection_name or self.default_collection_name
        metadata = metadata or {}

        try:
            # Ensure collection exists
            await self._ensure_collection_exists(collection_name)

            # Generate document ID
            doc_id = self._generate_document_id(content, metadata)

            # Check if document already exists
            if await self._document_exists(doc_id, collection_name):
                log_module_event(
                    "rag",
                    "document_exists",
                    {"document_id": doc_id, "collection": collection_name},
                )
                return doc_id

            # Chunk the document
            chunks = self._chunk_text(content)

            # Generate embeddings for all chunks in batch (more efficient)
            embeddings = await self._generate_embeddings(chunks, is_document=True)

            # Create document points
            points = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                aligned_embedding = self._align_embedding_dimension(
                    embedding, collection_name
                )
                chunk_id = str(uuid.uuid4())

                chunk_metadata = {
                    **metadata,
                    "document_id": doc_id,
                    "chunk_index": i,
                    "chunk_count": len(chunks),
                    "content": chunk,
                    "indexed_at": datetime.utcnow().isoformat(),
                }

                points.append(
                    PointStruct(
                        id=chunk_id, vector=aligned_embedding, payload=chunk_metadata
                    )
                )

            # Insert points into Qdrant
            self.qdrant_client.upsert(collection_name=collection_name, points=points)

            self.stats["documents_indexed"] += 1
            log_module_event(
                "rag",
                "document_indexed",
                {
                    "document_id": doc_id,
                    "collection": collection_name,
                    "chunks": len(chunks),
                    "metadata": metadata,
                },
            )

            return doc_id

        except Exception as e:
            logger.error(f"Error indexing document: {e}")
            log_module_event("rag", "indexing_failed", {"error": str(e)})
            raise

    async def index_processed_document(
        self, processed_doc: ProcessedDocument, collection_name: str = None
    ) -> str:
        """Index a processed document in the vector database"""
        if not self.enabled:
            raise RuntimeError("RAG module not initialized")

        collection_name = collection_name or self.default_collection_name

        try:
            # Special handling for JSONL files
            if processed_doc.file_type == "jsonl":
                # Import the optimized JSONL processor
                from app.services.jsonl_processor import JSONLProcessor

                jsonl_processor = JSONLProcessor(self)

                # Read the original file content
                with open(processed_doc.metadata.get("file_path", ""), "rb") as f:
                    file_content = f.read()

                # Process using the optimized JSONL processor
                return await jsonl_processor.process_and_index_jsonl(
                    collection_name=collection_name,
                    content=file_content,
                    filename=processed_doc.original_filename,
                    metadata=processed_doc.metadata,
                )

            # Ensure collection exists
            await self._ensure_collection_exists(collection_name)

            # Check if document already exists
            if await self._document_exists(processed_doc.id, collection_name):
                log_module_event(
                    "rag",
                    "document_exists",
                    {"document_id": processed_doc.id, "collection": collection_name},
                )
                return processed_doc.id

            # Chunk the document
            chunks = self._chunk_text(processed_doc.content)

            # Generate embeddings for all chunks in batch (more efficient)
            embeddings = await self._generate_embeddings(chunks, is_document=True)

            # Create document points with enhanced metadata
            points = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                aligned_embedding = self._align_embedding_dimension(
                    embedding, collection_name
                )
                chunk_id = str(uuid.uuid4())

                chunk_metadata = {
                    **processed_doc.metadata,
                    "document_id": processed_doc.id,
                    "original_filename": processed_doc.original_filename,
                    "file_type": processed_doc.file_type,
                    "mime_type": processed_doc.mime_type,
                    "language": processed_doc.language,
                    "entities": processed_doc.entities,
                    "keywords": processed_doc.keywords,
                    "word_count": processed_doc.word_count,
                    "sentence_count": processed_doc.sentence_count,
                    "file_hash": processed_doc.file_hash,
                    "processed_at": processed_doc.processed_at.isoformat(),
                    "chunk_index": i,
                    "chunk_count": len(chunks),
                    "content": chunk,
                    "indexed_at": datetime.utcnow().isoformat(),
                }

                # Add source_url if present in ProcessedDocument
                if processed_doc.source_url:
                    chunk_metadata["source_url"] = processed_doc.source_url

                points.append(
                    PointStruct(
                        id=chunk_id, vector=aligned_embedding, payload=chunk_metadata
                    )
                )

            # Insert points into Qdrant
            self.qdrant_client.upsert(collection_name=collection_name, points=points)

            self.stats["documents_indexed"] += 1
            log_module_event(
                "rag",
                "processed_document_indexed",
                {
                    "document_id": processed_doc.id,
                    "filename": processed_doc.original_filename,
                    "collection": collection_name,
                    "chunks": len(chunks),
                    "file_type": processed_doc.file_type,
                    "language": processed_doc.language,
                },
            )

            return processed_doc.id

        except Exception as e:
            logger.error(f"Error indexing processed document: {e}")
            log_module_event("rag", "indexing_failed", {"error": str(e)})
            raise

    async def _document_exists(
        self, document_id: str, collection_name: str = None
    ) -> bool:
        """Check if document exists in the collection"""
        collection_name = collection_name or self.default_collection_name

        try:
            result = self.qdrant_client.search(
                collection_name=collection_name,
                query_filter=Filter(
                    must=[
                        FieldCondition(
                            key="document_id", match=MatchValue(value=document_id)
                        )
                    ]
                ),
                limit=1,
            )
            return len(result) > 0
        except Exception:
            return False

    async def _hybrid_search(
        self,
        collection_name: str,
        query: str,
        query_vector: List[float],
        query_filter: Optional[Filter],
        limit: int,
        score_threshold: float,
    ) -> List[Any]:
        """Perform hybrid search combining vector similarity and BM25 scoring"""

        # Preprocess query for BM25
        query_terms = self._preprocess_text_for_bm25(query)

        # Get all documents from the collection (for BM25 scoring)
        # Note: In production, you'd want to optimize this with a proper BM25 index
        scroll_filter = query_filter or Filter()
        all_points = []

        # Use scroll to get all points
        offset = None
        batch_size = 100
        while True:
            search_result = self.qdrant_client.scroll(
                collection_name=collection_name,
                scroll_filter=scroll_filter,
                limit=batch_size,
                offset=offset,
                with_payload=True,
                with_vectors=False,
            )

            points = search_result[0]
            all_points.extend(points)

            if len(points) < batch_size:
                break

            offset = points[-1].id

        # Calculate BM25 scores for each document
        bm25_scores = {}
        for point in all_points:
            doc_id = point.payload.get("document_id", "")
            content = point.payload.get("content", "")

            # Calculate BM25 score
            bm25_score = self._calculate_bm25_score(query_terms, content)
            bm25_scores[doc_id] = bm25_score

        # Perform vector search
        vector_results = self.qdrant_client.search(
            collection_name=collection_name,
            query_vector=query_vector,
            query_filter=query_filter,
            limit=limit * 2,  # Get more results for re-ranking
            score_threshold=score_threshold / 2,  # Lower threshold for initial search
        )

        # Combine scores with improved normalization
        hybrid_weights = self.config.get("hybrid_weights", {"vector": 0.7, "bm25": 0.3})
        vector_weight = hybrid_weights.get("vector", 0.7)
        bm25_weight = hybrid_weights.get("bm25", 0.3)

        # Get score distributions for better normalization
        vector_scores = [r.score for r in vector_results]
        bm25_scores_list = list(bm25_scores.values())

        # Calculate statistics for normalization
        if vector_scores:
            v_max = max(vector_scores)
            v_min = min(vector_scores)
            v_range = v_max - v_min if v_max != v_min else 1
        else:
            v_max, v_min, v_range = 1, 0, 1

        if bm25_scores_list:
            bm25_max = max(bm25_scores_list)
            bm25_min = min(bm25_scores_list)
            bm25_range = bm25_max - bm25_min if bm25_max != bm25_min else 1
        else:
            bm25_max, bm25_min, bm25_range = 1, 0, 1

        # Create hybrid results with improved scoring
        hybrid_results = []
        for result in vector_results:
            doc_id = result.payload.get("document_id", "")
            vector_score = result.score
            bm25_score = bm25_scores.get(doc_id, 0.0)

            # Improved normalization using actual score distributions
            vector_norm = (vector_score - v_min) / v_range if v_range > 0 else 0.5
            bm25_norm = (bm25_score - bm25_min) / bm25_range if bm25_range > 0 else 0.5

            # Apply reciprocal rank fusion for better combination
            # This gives more weight to documents that rank highly in both methods
            rrf_vector = 1.0 / (
                1.0 + vector_results.index(result) + 1
            )  # +1 to avoid division by zero
            rrf_bm25 = (
                1.0
                / (1.0 + sorted(bm25_scores_list, reverse=True).index(bm25_score) + 1)
                if bm25_score in bm25_scores_list
                else 0
            )

            # Calculate hybrid score using both normalized scores and RRF
            hybrid_score = (
                vector_weight * vector_norm + bm25_weight * bm25_norm
            ) * 0.7 + (rrf_vector + rrf_bm25) * 0.3

            # Create new point with hybrid score
            hybrid_point = ScoredPoint(
                id=result.id,
                payload=result.payload,
                score=hybrid_score,
                vector=result.vector,
                shard_key=None,
                order_value=None,
            )
            hybrid_results.append(hybrid_point)

        # Sort by hybrid score and apply final threshold
        hybrid_results.sort(key=lambda x: x.score, reverse=True)
        final_results = [r for r in hybrid_results if r.score >= score_threshold][
            :limit
        ]

        logger.info(
            f"Hybrid search: {len(vector_results)} vector results, {len(final_results)} final results"
        )
        return final_results

    def _preprocess_text_for_bm25(self, text: str) -> List[str]:
        """Preprocess text for BM25 scoring"""
        if not NLTK_AVAILABLE:
            return text.lower().split()

        try:
            # Tokenize
            tokens = word_tokenize(text.lower())

            # Remove stopwords and non-alphabetic tokens
            stop_words = set(stopwords.words("english"))
            filtered_tokens = [
                token
                for token in tokens
                if token.isalpha() and token not in stop_words and len(token) > 2
            ]

            return filtered_tokens
        except:
            # Fallback to simple splitting
            return text.lower().split()

    def _calculate_bm25_score(self, query_terms: List[str], document: str) -> float:
        """Calculate BM25 score for a document against query terms"""
        if not query_terms:
            return 0.0

        # Preprocess document
        doc_terms = self._preprocess_text_for_bm25(document)
        if not doc_terms:
            return 0.0

        # Calculate term frequencies
        doc_len = len(doc_terms)
        avg_doc_len = 300  # Average document length (configurable)

        # BM25 parameters
        k1 = 1.2  # Controls term frequency saturation
        b = 0.75  # Controls document length normalization

        score = 0.0

        # Calculate IDF for each query term
        for term in set(query_terms):
            # Term frequency in document
            tf = doc_terms.count(term)

            # Simple IDF (log(N/n) + 1)
            # In production, you'd use the actual document frequency
            idf = 2.0  # Simplified IDF

            # BM25 formula
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * (doc_len / avg_doc_len))

            score += idf * (numerator / denominator)

        # Normalize score to 0-1 range
        return min(score / 10.0, 1.0)  # Simple normalization

    async def search_documents(
        self,
        query: str,
        max_results: int = None,
        filters: Dict[str, Any] = None,
        collection_name: str = None,
        score_threshold: float = None,
    ) -> List[SearchResult]:
        """Search for relevant documents"""
        if not self.enabled:
            raise RuntimeError("RAG module not initialized")

        collection_name = collection_name or self.default_collection_name
        max_results = max_results or self.config.get("max_results", 10)

        # Check cache (include collection name in cache key)
        cache_key = f"{collection_name}_{query}_{max_results}_{hash(str(filters))}"
        if cache_key in self.search_cache:
            self.stats["cache_hits"] += 1
            return self.search_cache[cache_key]

        try:
            import time

            start_time = time.time()

            # Generate query embedding with task-specific prefix for better retrieval
            optimized_query = f"query: {query}"
            query_embedding = await self._generate_embedding(optimized_query)
            query_embedding = self._align_embedding_dimension(
                query_embedding, collection_name
            )

            # Build filter
            search_filter = None
            if filters:
                conditions = []
                for key, value in filters.items():
                    conditions.append(
                        FieldCondition(key=key, match=MatchValue(value=value))
                    )
                search_filter = Filter(must=conditions)

            # Enhanced debugging for search
            logger.info("=== ENHANCED RAG SEARCH DEBUGGING ===")
            logger.info(f"Collection: {collection_name}")
            logger.info(f"Query: '{query}'")
            logger.info(f"Max results requested: {max_results}")
            logger.info(
                f"Query embedding (first 10 values): {query_embedding[:10] if query_embedding else 'None'}"
            )
            logger.info(
                f"Embedding service available: {self.embedding_service is not None}"
            )

            # Check if hybrid search is enabled
            enable_hybrid = self.config.get("enable_hybrid", False)
            # Use provided score_threshold or fall back to config
            search_score_threshold = (
                score_threshold
                if score_threshold is not None
                else self.config.get("score_threshold", 0.3)
            )

            if enable_hybrid and NLTK_AVAILABLE:
                # Perform hybrid search (vector + BM25)
                search_results = await self._hybrid_search(
                    collection_name=collection_name,
                    query=query,
                    query_vector=query_embedding,
                    query_filter=search_filter,
                    limit=max_results,
                    score_threshold=search_score_threshold,
                )
            else:
                # Pure vector search with improved threshold
                search_results = self.qdrant_client.search(
                    collection_name=collection_name,
                    query_vector=query_embedding,
                    query_filter=search_filter,
                    limit=max_results,
                    score_threshold=search_score_threshold,
                )

            logger.info(f"Raw search results count: {len(search_results)}")

            # Process results
            results = []
            document_scores = {}

            for i, result in enumerate(search_results):
                doc_id = result.payload.get("document_id")
                content = result.payload.get("content", "")
                score = result.score

                # Log each raw result for debugging
                logger.info(f"\n--- Raw Result {i+1} ---")
                logger.info(f"Score: {score}")
                logger.info(f"Document ID: {doc_id}")
                logger.info(f"Content preview (first 200 chars): {content[:200]}")
                logger.info(f"Metadata keys: {list(result.payload.keys())}")

                # Aggregate scores by document
                if doc_id in document_scores:
                    document_scores[doc_id]["score"] = max(
                        document_scores[doc_id]["score"], score
                    )
                    document_scores[doc_id]["content"] += "\n" + content
                else:
                    document_scores[doc_id] = {
                        "score": score,
                        "content": content,
                        "metadata": {
                            k: v
                            for k, v in result.payload.items()
                            if k not in ["content", "document_id"]
                        },
                    }

            logger.info(f"\nAggregated documents count: {len(document_scores)}")

            # Phase 2: URL Deduplication
            # Track documents by source_url to deduplicate
            url_to_doc = {}
            deduplicated_scores = {}
            docs_without_url = 0
            urls_deduplicated = 0

            for doc_id, data in document_scores.items():
                source_url = data["metadata"].get("source_url")

                if source_url:
                    # Document has a URL
                    if source_url in url_to_doc:
                        # URL already seen - keep document with higher score
                        existing_doc_id = url_to_doc[source_url]
                        existing_score = deduplicated_scores[existing_doc_id]["score"]

                        if data["score"] > existing_score:
                            # Replace with higher scoring document
                            logger.info(f"URL dedup: Replacing {existing_doc_id} (score={existing_score:.4f}) with {doc_id} (score={data['score']:.4f}) for URL: {source_url}")
                            del deduplicated_scores[existing_doc_id]
                            url_to_doc[source_url] = doc_id
                            deduplicated_scores[doc_id] = data
                        else:
                            logger.info(f"URL dedup: Skipping {doc_id} (score={data['score']:.4f}), keeping {existing_doc_id} (score={existing_score:.4f}) for URL: {source_url}")

                        urls_deduplicated += 1
                    else:
                        # First time seeing this URL
                        url_to_doc[source_url] = doc_id
                        deduplicated_scores[doc_id] = data
                else:
                    # Document without URL - always include
                    deduplicated_scores[doc_id] = data
                    docs_without_url += 1

            logger.info(f"\n=== URL Deduplication Metrics ===")
            logger.info(f"Documents before deduplication: {len(document_scores)}")
            logger.info(f"Documents after deduplication: {len(deduplicated_scores)}")
            logger.info(f"Unique URLs found: {len(url_to_doc)}")
            logger.info(f"Duplicate URLs removed: {urls_deduplicated}")
            logger.info(f"Documents without URL: {docs_without_url}")
            logger.info("=== END ENHANCED RAG SEARCH DEBUGGING ===")

            # Create SearchResult objects from deduplicated results
            for doc_id, data in deduplicated_scores.items():
                document = Document(
                    id=doc_id, content=data["content"], metadata=data["metadata"]
                )

                search_result = SearchResult(
                    document=document,
                    score=data["score"],
                    relevance_score=min(data["score"] * 100, 100),
                )

                results.append(search_result)

            # Sort by score
            results.sort(key=lambda x: x.score, reverse=True)

            # Update stats
            search_time = time.time() - start_time
            self.stats["searches_performed"] += 1
            self.stats["average_search_time"] = (
                self.stats["average_search_time"]
                * (self.stats["searches_performed"] - 1)
                + search_time
            ) / self.stats["searches_performed"]

            # Cache results
            self.search_cache[cache_key] = results

            log_module_event(
                "rag",
                "search_completed",
                {
                    "query": query,
                    "collection": collection_name,
                    "results_count": len(results),
                    "search_time": search_time,
                },
            )

            return results

        except Exception as e:
            logger.error(
                f"Error searching documents in collection {collection_name}: {e}"
            )
            log_module_event(
                "rag", "search_failed", {"error": str(e), "collection": collection_name}
            )
            raise

    async def delete_document(
        self, document_id: str, collection_name: str = None
    ) -> bool:
        """Delete a document from the vector database"""
        if not self.enabled:
            raise RuntimeError("RAG module not initialized")

        collection_name = collection_name or self.default_collection_name

        try:
            # Delete all chunks for this document
            self.qdrant_client.delete(
                collection_name=collection_name,
                points_selector=models.FilterSelector(
                    filter=Filter(
                        must=[
                            FieldCondition(
                                key="document_id", match=MatchValue(value=document_id)
                            )
                        ]
                    )
                ),
            )

            log_module_event(
                "rag",
                "document_deleted",
                {"document_id": document_id, "collection": collection_name},
            )
            return True

        except Exception as e:
            logger.error(
                f"Error deleting document from collection {collection_name}: {e}"
            )
            log_module_event(
                "rag",
                "deletion_failed",
                {"error": str(e), "collection": collection_name},
            )
            return False

    async def get_stats(self) -> Dict[str, Any]:
        """Get RAG module statistics"""
        stats = self.stats.copy()

        if self.enabled:
            try:
                # Use raw HTTP call to avoid Pydantic validation issues
                import httpx

                # Direct HTTP call to Qdrant API instead of using client to avoid Pydantic issues
                qdrant_url = f"http://{settings.QDRANT_HOST}:{settings.QDRANT_PORT}"

                async with httpx.AsyncClient() as client:
                    response = await client.get(
                        f"{qdrant_url}/collections/{self.default_collection_name}"
                    )

                    if response.status_code == 200:
                        collection_data = response.json()

                        # Safely extract stats from raw JSON
                        result = collection_data.get("result", {})

                        basic_stats = {
                            "total_points": result.get("points_count", 0),
                            "collection_status": result.get("status", "unknown"),
                        }

                        # Try to get vector dimension from config
                        try:
                            config = result.get("config", {})
                            params = config.get("params", {})
                            vectors = params.get("vectors", {})

                            if isinstance(vectors, dict) and "size" in vectors:
                                basic_stats["vector_dimension"] = vectors["size"]
                            else:
                                basic_stats["vector_dimension"] = "unknown"
                        except Exception as config_error:
                            logger.debug(
                                f"Could not get vector dimension: {config_error}"
                            )
                            basic_stats["vector_dimension"] = "unknown"

                        stats.update(basic_stats)
                    else:
                        # Collection doesn't exist or error
                        stats.update(
                            {
                                "total_points": 0,
                                "collection_status": "not_found",
                                "vector_dimension": "unknown",
                            }
                        )

            except Exception as e:
                logger.debug(f"Could not get Qdrant stats (using fallback): {e}")
                # Add basic fallback stats without logging as error since this is not critical
                stats.update(
                    {
                        "total_points": 0,
                        "collection_status": "unavailable",
                        "vector_dimension": "unknown",
                    }
                )
        else:
            stats.update(
                {
                    "total_points": 0,
                    "collection_status": "disabled",
                    "vector_dimension": "unknown",
                }
            )

        return stats

    async def process_request(
        self, request: Dict[str, Any], context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process a module request through the interceptor pattern"""
        if not self.enabled:
            raise RuntimeError("RAG module not initialized")

        action = request.get("action", "search")

        if action == "search":
            query = request.get("query")
            if not query:
                raise ValueError("Query is required for search action")

            max_results = request.get("max_results", self.config.get("max_results", 10))
            filters = request.get("filters", {})

            results = await self.search_documents(query, max_results, filters)

            return {
                "action": "search",
                "query": query,
                "results": [
                    {
                        "document_id": result.document.id,
                        "content": result.document.content,
                        "metadata": result.document.metadata,
                        "score": result.score,
                        "relevance_score": result.relevance_score,
                    }
                    for result in results
                ],
                "total_results": len(results),
                "cache_hit": False,  # Would be determined by search logic
            }

        elif action == "index":
            content = request.get("content")
            if not content:
                raise ValueError("Content is required for index action")

            metadata = request.get("metadata", {})
            document_id = await self.index_document(content, metadata)

            return {
                "action": "index",
                "document_id": document_id,
                "status": "success",
                "message": "Document indexed successfully",
            }

        elif action == "process":
            file_data = request.get("file_data")
            filename = request.get("filename")
            if not file_data or not filename:
                raise ValueError(
                    "File data and filename are required for process action"
                )

            # Decode base64 file data if provided as string
            if isinstance(file_data, str):
                import base64

                file_data = base64.b64decode(file_data)

            metadata = request.get("metadata", {})
            processed_doc = await self.process_document(file_data, filename, metadata)

            return {
                "action": "process",
                "document_id": processed_doc.id,
                "filename": processed_doc.original_filename,
                "file_type": processed_doc.file_type,
                "mime_type": processed_doc.mime_type,
                "word_count": processed_doc.word_count,
                "sentence_count": processed_doc.sentence_count,
                "language": processed_doc.language,
                "entities_count": len(processed_doc.entities),
                "keywords_count": len(processed_doc.keywords),
                "processing_time": processed_doc.processing_time,
                "status": "success",
                "message": "Document processed successfully",
            }

        elif action == "delete":
            document_id = request.get("document_id")
            if not document_id:
                raise ValueError("Document ID is required for delete action")

            success = await self.delete_document(document_id)

            return {
                "action": "delete",
                "document_id": document_id,
                "status": "success" if success else "failed",
                "message": "Document deleted successfully"
                if success
                else "Failed to delete document",
            }

        elif action == "stats":
            stats = await self.get_stats()

            return {"action": "stats", "statistics": stats}

        else:
            raise ValueError(f"Unsupported action: {action}")

    async def pre_request_interceptor(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Pre-request interceptor for RAG enhancement"""
        if not self.enabled:
            return context

        request = context.get("request")
        if not request:
            return context

        # Check if this is a request that could benefit from RAG
        if request.url.path.startswith("/api/v1/chat") or request.url.path.startswith(
            "/api/v1/completions"
        ):
            # Extract query/prompt from request
            request_body = await request.body() if hasattr(request, "body") else b""

            if request_body:
                try:
                    data = json.loads(request_body.decode())
                    query = data.get("message", data.get("prompt", ""))

                    if query:
                        # Search for relevant documents
                        search_results = await self.search_documents(
                            query, max_results=3
                        )

                        if search_results:
                            # Add context to request
                            context["rag_context"] = [
                                {
                                    "content": result.document.content,
                                    "metadata": result.document.metadata,
                                    "relevance_score": result.relevance_score,
                                }
                                for result in search_results
                            ]

                            log_module_event(
                                "rag",
                                "context_added",
                                {
                                    "query": query[:100],
                                    "results_count": len(search_results),
                                },
                            )

                except Exception as e:
                    logger.error(f"Error processing RAG request: {e}")

        return context


# Global RAG instance
rag_module = RAGModule()


# Module interface functions
async def initialize(config: Dict[str, Any]):
    """Initialize RAG module"""
    await rag_module.initialize(config)


async def cleanup():
    """Cleanup RAG module"""
    await rag_module.cleanup()


async def pre_request_interceptor(context: Dict[str, Any]) -> Dict[str, Any]:
    """Pre-request interceptor"""
    return await rag_module.pre_request_interceptor(context)


# Additional exported functions
async def process_document(
    file_data: bytes, filename: str, metadata: Dict[str, Any] = None
) -> ProcessedDocument:
    """Process a document with full content analysis"""
    return await rag_module.process_document(file_data, filename, metadata)


async def index_document(
    content: str, metadata: Dict[str, Any] = None, collection_name: str = None
) -> str:
    """Index a document (backward compatibility)"""
    return await rag_module.index_document(content, metadata, collection_name)


async def index_processed_document(
    processed_doc: ProcessedDocument, collection_name: str = None
) -> str:
    """Index a processed document"""
    return await rag_module.index_processed_document(processed_doc, collection_name)


async def search_documents(
    query: str,
    max_results: int = None,
    filters: Dict[str, Any] = None,
    collection_name: str = None,
    score_threshold: float = None,
) -> List[SearchResult]:
    """Search documents"""
    return await rag_module.search_documents(
        query, max_results, filters, collection_name, score_threshold
    )


async def delete_document(document_id: str, collection_name: str = None) -> bool:
    """Delete a document"""
    return await rag_module.delete_document(document_id, collection_name)


async def create_collection(collection_name: str) -> bool:
    """Create a new Qdrant collection"""
    return await rag_module.create_collection(collection_name)


async def delete_collection(collection_name: str) -> bool:
    """Delete a Qdrant collection"""
    return await rag_module.delete_collection(collection_name)


async def get_supported_types() -> List[str]:
    """Get list of supported file types"""
    return list(rag_module.supported_types.keys())
